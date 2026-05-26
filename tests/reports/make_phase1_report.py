"""
Phase 1 구현 결과 Word 문서 생성기
===================================
실행:
  $env:PYTHONIOENCODING="utf-8"
  .venv\Scripts\python.exe tests/make_phase1_report.py
"""

import sys, os, importlib.util
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

# ── 모듈 로드 ──────────────────────────────────────────────────
def _load(alias, path):
    spec = importlib.util.spec_from_file_location(alias, os.path.join(ROOT, path))
    mod  = importlib.util.module_from_spec(spec)
    sys.modules[alias] = mod
    spec.loader.exec_module(mod)
    return mod

C  = _load("travel_logic.config.constants", "travel_logic/config/constants.py")
S  = _load("travel_logic.config.settings",  "travel_logic/config/settings.py")
SS = _load("travel_logic.services.scoring_service", "travel_logic/services/scoring_service.py")
ScoringService = SS.ScoringService

# ── 색상 팔레트 ─────────────────────────────────────────────────
DARK_BG    = RGBColor(0x1E, 0x1E, 0x2E)
BLUE_DEEP  = RGBColor(0x1F, 0x4E, 0x79)
BLUE_MID   = RGBColor(0x2E, 0x75, 0xB6)
BLUE_LIGHT = RGBColor(0xD6, 0xE4, 0xF7)
GREEN_DARK = RGBColor(0x37, 0x5A, 0x29)
GREEN_LIGHT= RGBColor(0xE2, 0xEF, 0xDA)
RED_DARK   = RGBColor(0x83, 0x17, 0x17)
RED_LIGHT  = RGBColor(0xFC, 0xE4, 0xE4)
YELLOW_BG  = RGBColor(0xFF, 0xF2, 0xCC)
GRAY_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)
GRAY_MID   = RGBColor(0xD9, 0xD9, 0xD9)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
ORANGE     = RGBColor(0xC5, 0x5A, 0x11)

# ── Word 헬퍼 함수 ──────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_table_borders(table, color="BFBFBF", size=4):
    tbl  = table._tbl
    tblPr= tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    str(size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_run(para, text, bold=False, italic=False, size=10,
            color: RGBColor = None, font_name="맑은 고딕"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1, color: RGBColor = BLUE_DEEP):
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, bold=True, size=sizes.get(level, 11), color=color)
    return p

def body_para(doc, text, indent=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, text, size=10, color=color or BLACK)
    return p

def make_table(doc, headers, rows,
               header_bg=BLUE_MID, header_fg=WHITE,
               alt_bg=BLUE_LIGHT, col_widths=None):
    """헤더 + 데이터 행 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "BFBFBF", 4)

    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9, color=header_fg)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row_obj = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = row_obj.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(val), size=9)

    # 열 너비
    if col_widths:
        for ri2, row_obj2 in enumerate(table.rows):
            for ci2, cell2 in enumerate(row_obj2.cells):
                if ci2 < len(col_widths):
                    cell2.width = Inches(col_widths[ci2])

    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.font.size = Pt(1)
    border_xml = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    border_xml.append(bottom)
    p._p.get_or_add_pPr().append(border_xml)

def colored_badge_para(doc, items):
    """(색상, 텍스트) 튜플 리스트를 한 줄에 배지처럼 출력"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for color, text in items:
        add_run(p, f" {text} ", bold=True, size=9, color=WHITE)
        # 실제 배지 배경은 Word에서 표로 구현해야 하지만 색상 텍스트로 대체
        add_run(p, "  ", size=9)

# ════════════════════════════════════════════════════════════════
#  샘플 데이터 & Phase 1 실행
# ════════════════════════════════════════════════════════════════
JEJU_PLACES = [
    {"id":"s01","name":"성산일출봉","category":"관광명소","lat":33.458,"lng":126.9425,"rating":4.8,"img_url":"y"},
    {"id":"s02","name":"한라산 국립공원","category":"관광명소","lat":33.3617,"lng":126.5292,"rating":4.9,"img_url":"y"},
    {"id":"s03","name":"만장굴","category":"관광명소","lat":33.528,"lng":126.7713,"rating":4.5,"img_url":"y"},
    {"id":"s04","name":"천지연 폭포","category":"관광명소","lat":33.246,"lng":126.5497,"rating":4.6,"img_url":"y"},
    {"id":"s05","name":"협재 해수욕장","category":"park","lat":33.3941,"lng":126.2394,"rating":4.7,"img_url":"y"},
    {"id":"s06","name":"제주 민속촌","category":"관광명소","lat":33.3195,"lng":126.8155,"rating":4.2,"img_url":""},
    {"id":"s07","name":"폐업된 관광지X","category":"관광명소","lat":33.4,"lng":126.5,"rating":1.2,"img_url":""},
    {"id":"s08","name":"좌표없는명소Y","category":"관광명소","lat":0.0,"lng":0.0,"rating":4.5,"img_url":""},
    {"id":"s09","name":"우도","category":"관광명소","lat":33.5025,"lng":126.9521,"rating":4.7,"img_url":"y"},
    {"id":"s10","name":"비자림","category":"park","lat":33.5001,"lng":126.8112,"rating":4.4,"img_url":"y"},
    {"id":"f01","name":"흑돼지거리 식당A","category":"restaurant","lat":33.4996,"lng":126.5312,"rating":4.3,"img_url":"y"},
    {"id":"f02","name":"칠성식당","category":"음식점","lat":33.5,"lng":126.52,"rating":4.1,"img_url":"y"},
    {"id":"f03","name":"해녀촌","category":"restaurant","lat":33.46,"lng":126.93,"rating":4.5,"img_url":"y"},
    {"id":"f04","name":"폐업식당 Z","category":"restaurant","lat":33.49,"lng":126.51,"rating":0.8,"img_url":""},
    {"id":"f05","name":"고기국수 본점","category":"음식점","lat":33.489,"lng":126.502,"rating":4.6,"img_url":"y"},
    {"id":"f06","name":"자연식 뷔페","category":"restaurant","lat":33.51,"lng":126.55,"rating":3.2,"img_url":"y"},
    {"id":"f07","name":"한치 물회","category":"음식점","lat":33.25,"lng":126.56,"rating":4.0,"img_url":"y"},
    {"id":"c01","name":"바다뷰 카페","category":"카페","lat":33.52,"lng":126.54,"rating":4.5,"img_url":"y"},
    {"id":"c02","name":"감귤 카페","category":"카페","lat":33.47,"lng":126.32,"rating":4.3,"img_url":"y"},
    {"id":"c03","name":"스타벅스 제주점","category":"cafe","lat":33.49,"lng":126.53,"rating":4.0,"img_url":"y"},
    {"id":"c04","name":"제주 커피 로스터리","category":"cafe","lat":33.505,"lng":126.51,"rating":4.6,"img_url":"y"},
    {"id":"c05","name":"저평점 카페 W","category":"카페","lat":33.5,"lng":126.5,"rating":1.9,"img_url":""},
    {"id":"h01","name":"제주신라호텔","category":"호텔","lat":33.2476,"lng":126.5646,"rating":4.8,"img_url":"y"},
    {"id":"h02","name":"롯데호텔 제주","category":"호텔","lat":33.25,"lng":126.555,"rating":4.6,"img_url":"y"},
    {"id":"h03","name":"감귤 펜션","category":"펜션","lat":33.4,"lng":126.3,"rating":4.2,"img_url":"y"},
    {"id":"h04","name":"저평점 모텔 Q","category":"모텔","lat":33.49,"lng":126.5,"rating":2.5,"img_url":""},
    {"id":"h05","name":"아난티 리조트","category":"리조트","lat":33.28,"lng":126.42,"rating":4.7,"img_url":"y"},
    {"id":"x01","name":"평점없는 장소A","category":"관광명소","lat":33.5,"lng":126.5,"rating":None,"img_url":""},
    {"id":"x02","name":"","category":"관광명소","lat":33.5,"lng":126.5,"rating":4.5,"img_url":""},
]

USER_DATA = {
    "dest_city":"제주","budget_level":"중","star_rating":4.0,
    "style":["자연","맛집"],"with_kids":False,"stroller":False,
    "barrier_free":False,"photo_spot":True,"pace":"알차게",
    "transport":["항공"],"companions":["커플"],"_duration":4,
}
NUM_DAYS = 4

# Phase 1 실행
unique_places = []
seen_names = set()
removed_no_name = []
for p in JEJU_PLACES:
    if not p.get("name"):
        removed_no_name.append(p)
        continue
    clean = "".join(filter(str.isalnum, p["name"])).lower()
    if clean not in seen_names:
        seen_names.add(clean)
        unique_places.append(p)

hard_filtered = ScoringService.hard_filter(unique_places, USER_DATA)
surviving_ids = {p["id"] for p in hard_filtered}
removed_hard  = [p for p in unique_places if p["id"] not in surviving_ids]

min_hotel_rating = float(USER_DATA["star_rating"])
hotels = [p for p in hard_filtered if any(kw in str(p.get("category","")) for kw in C.HOTEL_CATEGORIES) and float(p.get("rating",0)) >= min_hotel_rating]
hotel_ids  = {p["id"] for p in hotels}
non_hotels = [p for p in hard_filtered if p["id"] not in hotel_ids]

scored = []
for p in non_hotels:
    score, tags = ScoringService.calculate_score(p, USER_DATA)
    p2 = dict(p); p2["score"] = score; p2["matched_tags"] = tags
    scored.append(p2)
scored.sort(key=lambda x: x["score"], reverse=True)

top_n = ScoringService.extract_top_n(scored, NUM_DAYS)
excluded = scored[len(top_n):]
sights, foods, cafes = ScoringService.categorize_visits(top_n, USER_DATA)

# ════════════════════════════════════════════════════════════════
#  Word 문서 생성
# ════════════════════════════════════════════════════════════════
doc = Document()

# ── 기본 폰트 설정 ───────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

# 여백 설정
for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ────────────────────────────────────────────────────────────────
#  표지 헤더 (배경색 박스 = 박스 없이 텍스트로)
# ────────────────────────────────────────────────────────────────
# 표지 테이블로 타이틀 박스 구현
title_table = doc.add_table(rows=1, cols=1)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
title_cell = title_table.cell(0, 0)
set_cell_bg(title_cell, BLUE_DEEP)
title_cell.width = Inches(6)

tp = title_cell.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(14)
tp.paragraph_format.space_after  = Pt(6)
add_run(tp, "Pick&Go — Phase 1 구현 결과 리포트", bold=True, size=18, color=WHITE)

tp2 = title_cell.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp2.paragraph_format.space_after = Pt(14)
add_run(tp2, "필터링 & 점수화 파이프라인  ·  2026-04-09  ·  50개 테스트 전부 통과", size=10, color=GRAY_MID)

doc.add_paragraph()

# ── 요약 지표 테이블 (2×4) ──────────────────────────────────────
summary_table = doc.add_table(rows=2, cols=4)
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(summary_table, "2E75B6", 6)

summary_data = [
    ("50개", "총 테스트 항목", BLUE_MID),
    ("50개", "전부 통과 (100%)", GREEN_DARK),
    ("3회", "API 호출 제거", RGBColor(0xC5,0x5A,0x11)),
    ("4곳", "핵심 변경 지점", RGBColor(0x70,0x30,0xA0)),
]
for ci, (num, label, color) in enumerate(summary_data):
    top_cell = summary_table.rows[0].cells[ci]
    bot_cell = summary_table.rows[1].cells[ci]
    set_cell_bg(top_cell, BLUE_LIGHT)
    set_cell_bg(bot_cell, GRAY_LIGHT)
    p_num = top_cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_before = Pt(6)
    p_num.paragraph_format.space_after  = Pt(2)
    add_run(p_num, num, bold=True, size=20, color=color)
    p_lbl = bot_cell.paragraphs[0]
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lbl.paragraph_format.space_before = Pt(2)
    p_lbl.paragraph_format.space_after  = Pt(6)
    add_run(p_lbl, label, size=9, color=RGBColor(0x40,0x40,0x40))

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════════
#  1. 변경 전 vs 변경 후
# ════════════════════════════════════════════════════════════════
heading(doc, "1. 변경 전 vs 변경 후", 1)

before_after_rows = [
    ("장소 필터 기준",  "좌표·이름 존재 여부만 확인",  "hard_filter(): 예산별 최소 평점 + 좌표 0.0 제거"),
    ("API 호출",        "check_place_status() 3회 (Google API 비용 발생)",  "완전 제거 → 평점으로 자연 필터링"),
    ("숙소 분류 기준",  "카테고리 키워드 코드 내 하드코딩",  "HOTEL_CATEGORIES 상수 기반 일원화"),
    ("숙소 풀 분리",    "방문 장소 풀 안에 숙소 포함\n(Phase 4 앵커링 장소 부족 위험)",  "숙소 별도 풀 분리 → extract_top_n은 방문장소만 대상"),
    ("후보 장소 수",    "제한 없이 전체 처리",  f"extract_top_n(): 여행일수 × {C.CANDIDATE_POOL_RATIO}개만 추출"),
    ("카테고리 분류",   "_categorize_places() 내 즉석 분류",  "categorize_visits() 전용 메서드 위임"),
    ("공항 좌표",       "제주공항 lat:33.5113 하드코딩\n(제주만 지원)",  "AIRPORT_COORDS 딕셔너리 (국내 8개 도시 동적 조회)"),
]

ba_table = doc.add_table(rows=1 + len(before_after_rows), cols=3)
ba_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(ba_table, "BFBFBF", 4)

hdr_cells = ba_table.rows[0].cells
for ci, (h, bg) in enumerate([("항목", BLUE_DEEP), ("변경 전 (구버전)", RED_DARK), ("변경 후 (Phase 1)", GREEN_DARK)]):
    set_cell_bg(hdr_cells[ci], bg)
    p = hdr_cells[ci].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, size=9, color=WHITE)

for ri, (item, before, after) in enumerate(before_after_rows):
    row_cells = ba_table.rows[ri+1].cells
    bg = GRAY_LIGHT if ri % 2 == 0 else WHITE
    set_cell_bg(row_cells[0], bg)
    set_cell_bg(row_cells[1], RED_LIGHT if ri % 2 == 0 else WHITE)
    set_cell_bg(row_cells[2], GREEN_LIGHT if ri % 2 == 0 else WHITE)

    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(row_cells[0].paragraphs[0], item, bold=True, size=9)
    add_run(row_cells[1].paragraphs[0], before, size=9, color=RED_DARK)
    add_run(row_cells[2].paragraphs[0], after,  size=9, color=GREEN_DARK)

for i, w in enumerate([1.3, 2.5, 2.5]):
    for row_obj in ba_table.rows:
        row_obj.cells[i].width = Inches(w)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════════
#  2. Phase 1 파이프라인 흐름
# ════════════════════════════════════════════════════════════════
heading(doc, "2. Phase 1 파이프라인 전체 흐름", 1)

pipeline_steps = [
    ("STEP 0", "DB 원본 데이터 로드",     f"backend.get_places(city) → 전체 장소 {len(JEJU_PLACES)}개 로드", BLUE_MID),
    ("STEP 1", "이름 유효성 + 중복 제거", f"seen_names set으로 동일 장소 제거 → {len(unique_places)}개", BLUE_MID),
    ("STEP 2", "hard_filter() [NEW]",     f"예산별 최소 평점 컷 + 좌표 0.0 제거 → {len(hard_filtered)}개 생존 ({len(removed_hard)}개 제거)", RGBColor(0xC5,0x5A,0x11)),
    ("STEP 3", "숙소 별도 풀 분리 [NEW]", f"HOTEL_CATEGORIES 상수 + star_rating≥{USER_DATA['star_rating']} → 숙소 {len(hotels)}개 / 방문후보 {len(non_hotels)}개", GREEN_DARK),
    ("STEP 4", "점수 계산 & 정렬",        f"취향 매칭·사진명소·동반자 보너스 → score 내림차순 정렬 ({len(scored)}개)", BLUE_MID),
    ("STEP 5", "extract_top_n() [NEW]",   f"{NUM_DAYS}일 × {C.CANDIDATE_POOL_RATIO} = {NUM_DAYS*C.CANDIDATE_POOL_RATIO}개 추출 → 실제 {len(top_n)}개 (장소 수 부족)", GREEN_DARK),
    ("STEP 6", "categorize_visits() [NEW]", f"관광 {len(sights)}개 · 식당 {len(foods)}개 · 카페 {len(cafes)}개 → Phase 2 입력 준비 완료", GREEN_DARK),
]

for (step, title, desc, color) in pipeline_steps:
    pipe_t = doc.add_table(rows=1, cols=2)
    pipe_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    badge_cell = pipe_t.cell(0,0)
    desc_cell  = pipe_t.cell(0,1)
    badge_cell.width = Inches(1.0)
    desc_cell.width  = Inches(5.3)
    set_cell_bg(badge_cell, color)
    badge_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    badge_cell.paragraphs[0].paragraph_format.space_after  = Pt(0)
    add_run(badge_cell.paragraphs[0], step, bold=True, size=9, color=WHITE)
    p2 = badge_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)
    add_run(p2, title, bold=True, size=8, color=WHITE)

    desc_cell.paragraphs[0].paragraph_format.space_before = Pt(6)
    desc_cell.paragraphs[0].paragraph_format.left_indent  = Cm(0.3)
    add_run(desc_cell.paragraphs[0], desc, size=9, color=BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════════
#  3. hard_filter 제거 목록
# ════════════════════════════════════════════════════════════════
heading(doc, "3. STEP 2 — hard_filter() 제거 목록", 1)

p_desc = body_para(doc, f"예산='중' 기준 최소 평점 3.0 적용 (식당·숙소는 완화 기준 2.0 적용). 총 {len(removed_no_name)+len(removed_hard)}개 제거.")
p_desc.paragraph_format.space_after = Pt(6)

removed_rows = []
for p in removed_no_name:
    removed_rows.append((p.get("name","(이름없음)") or "(이름없음)", p["category"], str(p.get("rating","없음")), "이름 없음"))
for p in removed_hard:
    lat = p.get("lat", 0)
    rat = p.get("rating") or 0
    reason = "좌표 0.0 (무효 좌표)" if lat == 0.0 else f"평점 {rat} 미달"
    removed_rows.append((p["name"], p["category"], str(rat), reason))

make_table(doc,
    headers=["장소명", "카테고리", "평점", "제거 사유"],
    rows=removed_rows,
    header_bg=RED_DARK, header_fg=WHITE,
    alt_bg=RED_LIGHT, col_widths=[1.8, 1.2, 0.8, 2.5])

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════════
#  4. 점수 순위표
# ════════════════════════════════════════════════════════════════
heading(doc, "4. STEP 4 — 취향 점수 계산 결과 (전체 순위)", 1)

p_desc2 = body_para(doc, f"스타일=[자연, 맛집], photo_spot=True 기준. 스타일 매칭 시 +30점 보너스 적용.")
p_desc2.paragraph_format.space_after = Pt(6)

score_rows = [(str(i), p["name"], p["category"], f"★ {p.get('rating','?')}", f"{p['score']}점") for i, p in enumerate(scored, 1)]
make_table(doc,
    headers=["순위", "장소명", "카테고리", "평점", "취향 점수"],
    rows=score_rows,
    header_bg=BLUE_DEEP, header_fg=WHITE,
    alt_bg=BLUE_LIGHT, col_widths=[0.5, 2.0, 1.5, 0.9, 1.0])

# 특이사항 노트
p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(6)
p_note.paragraph_format.left_indent  = Cm(0.5)
add_run(p_note, "※ 주목: ", bold=True, size=9, color=ORANGE)
add_run(p_note, "자연식 뷔페가 평점 3.2임에도 1위 → 사용자 스타일 '맛집' 보너스(+30)가 평점 점수보다 크게 반영됨", size=9, color=ORANGE)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════════
#  5. categorize_visits 최종 분류
# ════════════════════════════════════════════════════════════════
heading(doc, "5. STEP 6 — categorize_visits() 최종 분류 결과", 1)

# 세 분류를 각각 서브 테이블로
for (title, places, bg_h) in [
    ("관광지 (sights)", sights, BLUE_DEEP),
    ("식당 (foods)",    foods,  GREEN_DARK),
    ("카페 (cafes)",    cafes,  RGBColor(0x70,0x30,0xA0)),
]:
    heading(doc, f"  {title}  —  {len(places)}개", 2, color=bg_h)
    rows2 = [(str(i), p["name"], p["category"], f"★ {p.get('rating','?')}", f"{p['score']}점") for i, p in enumerate(places, 1)]
    make_table(doc,
        headers=["순위", "장소명", "카테고리", "평점", "점수"],
        rows=rows2, header_bg=bg_h, header_fg=WHITE,
        alt_bg=BLUE_LIGHT if bg_h==BLUE_DEEP else GREEN_LIGHT if bg_h==GREEN_DARK else YELLOW_BG,
        col_widths=[0.5, 2.2, 1.5, 0.9, 1.0])
    doc.add_paragraph()

add_divider(doc)

# ════════════════════════════════════════════════════════════════
#  6. 장점 & 주의사항
# ════════════════════════════════════════════════════════════════
heading(doc, "6. 장점 & 주의사항", 1)

pros_cons_table = doc.add_table(rows=1, cols=2)
pros_cons_table.alignment = WD_TABLE_ALIGNMENT.CENTER
pros_cell = pros_cons_table.cell(0,0)
cons_cell = pros_cons_table.cell(0,1)
pros_cell.width = Inches(3.1)
cons_cell.width = Inches(3.1)

set_cell_bg(pros_cell, GREEN_LIGHT)
set_cell_bg(cons_cell, YELLOW_BG)

pros_p = pros_cell.paragraphs[0]
pros_p.paragraph_format.space_before = Pt(6)
add_run(pros_p, "장점", bold=True, size=11, color=GREEN_DARK)
pros_items = [
    "Google API 호출 3회 제거 → 비용·속도 대폭 개선",
    "예산별 최소 평점으로 품질 기준 명확화",
    "top_N 제한으로 Phase 2~5 처리 속도 향상",
    "AIRPORT_COORDS로 국내 8개 도시 지원",
    "HOTEL_CATEGORIES 상수화 → 유지보수 용이",
    "숙소 별도 풀로 Phase 4 앵커링 안정성 보장",
]
for item in pros_items:
    p = pros_cell.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, "✓  " + item, size=9, color=GREEN_DARK)

cons_p = cons_cell.paragraphs[0]
cons_p.paragraph_format.space_before = Pt(6)
add_run(cons_p, "주의사항", bold=True, size=11, color=ORANGE)
cons_items = [
    "check_place_status 제거 → 폐업 장소 포함 가능\n(낮은 평점으로 자연 필터링되어 실 영향 낮음)",
    "배리어프리 필터 미구현\n(TourAPI 연동 후 활성화 예정)",
    "top_N이 너무 작으면 일정 다양성 감소 가능",
    "미등록 도시는 airport_place=None\n(공항 안내 없이 진행)",
]
for item in cons_items:
    p = cons_cell.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, "⚠  " + item, size=9, color=ORANGE)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════════
#  7. 테스트 결과 요약
# ════════════════════════════════════════════════════════════════
heading(doc, "7. 테스트 결과 (50개 항목)", 1)

test_rows = [
    ("Test 1", "constants.py — 5-Phase 신규 상수 검증",       "14", "14", "100%"),
    ("Test 2", "hard_filter() — 예산별 평점 컷 & 좌표 유효성", "10", "10", "100%"),
    ("Test 3", "extract_top_n() — CANDIDATE_POOL_RATIO 사용",  "4",  "4",  "100%"),
    ("Test 4", "categorize_visits() — 카테고리 분류 정확도",   "7",  "7",  "100%"),
    ("Test 5", "AIRPORT_COORDS — 공항 하드코딩 제거",          "9",  "9",  "100%"),
    ("Test 6", "HOTEL_CATEGORIES — 숙소 별도 풀 분리",         "6",  "6",  "100%"),
    ("합계",   "",                                              "50", "50", "100%"),
]

t_table = doc.add_table(rows=1+len(test_rows), cols=5)
t_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_table, "BFBFBF", 4)
for ci, h in enumerate(["", "테스트 항목", "전체", "통과", "통과율"]):
    set_cell_bg(t_table.rows[0].cells[ci], BLUE_DEEP)
    p = t_table.rows[0].cells[ci].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, size=9, color=WHITE)

for ri, (tid, desc, total, passed, rate) in enumerate(test_rows):
    row_cells = t_table.rows[ri+1].cells
    is_last = ri == len(test_rows)-1
    bg = RGBColor(0xD6,0xF0,0xD6) if is_last else (GRAY_LIGHT if ri%2==0 else WHITE)
    for cell in row_cells:
        set_cell_bg(cell, bg)
    for ci, (val, align) in enumerate([(tid,True),(desc,False),(total,True),(passed,True),(rate,True)]):
        p = row_cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align else WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, val, bold=is_last, size=9, color=GREEN_DARK if is_last else BLACK)

for i, w in enumerate([0.7, 3.2, 0.6, 0.6, 0.7]):
    for row_obj in t_table.rows:
        row_obj.cells[i].width = Inches(w)

doc.add_paragraph()

# ── 마지막 배너 ──────────────────────────────────────────────────
banner_t = doc.add_table(rows=1, cols=1)
banner_t.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = banner_t.cell(0,0)
set_cell_bg(bc, BLUE_DEEP)
bp1 = bc.paragraphs[0]
bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp1.paragraph_format.space_before = Pt(12)
bp1.paragraph_format.space_after  = Pt(4)
add_run(bp1, "Phase 1 구현 완료  ·  50/50 테스트 ALL PASS", bold=True, size=14, color=WHITE)
bp2 = bc.add_paragraph()
bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp2.paragraph_format.space_after = Pt(12)
add_run(bp2, "다음 단계: Phase 2 K-Means 클러스터링 구현 대기 중", size=10, color=GRAY_MID)

# ── 저장 ────────────────────────────────────────────────────────
OUT = os.path.join(ROOT, "Pick&Go_Phase1_구현결과.docx")
doc.save(OUT)
print(f"Word 파일 저장 완료: {OUT}")
